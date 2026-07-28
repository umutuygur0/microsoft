import io

import pytest

from src.document_readers import extract_docx_text, extract_pdf_text, extract_text


def _build_docx_bytes(paragraphs: list[str]) -> bytes:
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_blank_pdf_bytes() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_extract_docx_text_reads_paragraphs():
    docx_bytes = _build_docx_bytes(["First paragraph.", "Second paragraph about gas leaks."])
    text = extract_docx_text(docx_bytes)
    assert "First paragraph." in text
    assert "Second paragraph about gas leaks." in text


def test_extract_docx_text_returns_empty_for_garbage_bytes():
    assert extract_docx_text(b"not a real docx file") == ""


def test_extract_docx_text_converts_heading_style_to_markdown_h1():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Freeze Protocol", style="Heading 1")
    doc.add_paragraph("Isolate the line and notify the control room.")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_docx_text(buf.getvalue())
    assert "# Freeze Protocol" in text.splitlines()


def test_extract_pdf_text_returns_empty_for_blank_page():
    # A structurally valid PDF with no text layer -> empty string, not a crash.
    assert extract_pdf_text(_build_blank_pdf_bytes()) == ""


def test_extract_pdf_text_returns_empty_for_garbage_bytes():
    assert extract_pdf_text(b"not a real pdf file") == ""


def test_extract_text_dispatches_pdf():
    assert extract_text("report.pdf", _build_blank_pdf_bytes()) == ""


def test_extract_text_dispatches_docx():
    text = extract_text("notes.docx", _build_docx_bytes(["Hello world."]))
    assert "Hello world." in text


def test_extract_text_raises_for_unsupported_extension():
    with pytest.raises(ValueError):
        extract_text("script.py", b"print('hi')")
